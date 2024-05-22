import docx
import docx.shared
import openpyxl
from module import function, Branches

doc = docx.Document()
doc.add_paragraph('Практическая работа №5')
doc.add_paragraph('Задание: Разработать программу для табулирования кусочно-ломанной функции')
doc.add_picture('screenshot.jpg', width=docx.shared.Cm(10))
table = doc.add_table(rows=1, cols=2)
heading_row = table.rows[0].cells
heading_row[0].text = 'X'
heading_row[1].text = 'Y'

workbook = openpyxl.Workbook()
sheet = workbook.active

branching = Branches(0, 0, 0, 0, 0)

try:
    a = float(input('Введите начало отрезка: '))
    b = float(input('Введите конец отрезка: '))
    h = float(input('Введите шаг: '))
    branching.input()
except:
    print("Произошла ошибка ввода")
    exit(0)

n = int((b - a) // h + 1)

x = a

for i in range(n):
    y = function(i, branching)
    print(f'х = {x} \ty = {y}')

    cell = sheet.cell(row=i+1, column=1)
    cell.value = x
    cell = sheet.cell(row=i+1, column=2)
    cell.value = y

    cells = table.add_row().cells
    cells[0].text = str(x)
    cells[1].text = str(round(y, 3))

    x += h

doc.save('test.docx')
workbook.save('Tabulirovanie.xlsx')

print('\nЧтение таблицы из файла')

workbook = openpyxl.load_workbook('Tabulirovanie.xlsx')
worksheet = workbook.active
for i in range(worksheet.max_row):
    for col in worksheet.iter_cols(1, worksheet.max_column):
        print(float(col[i].value), end=' \t\t')
    print('')



print('\nКонец')